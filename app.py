from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Cm
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, KeepTogether
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PIL import Image as PILImage
import io
import os
import re

app = Flask(__name__)

def replace_placeholders(doc, fields):
    aliases = {
        'doc_date_day': 'contract_date_day',
        'doc_date_month': 'contract_date_month',
        'doc_number': 'contract_number',
        'doc_date': 'contract_date',
    }
    
    extended_fields = dict(fields)
    for original, alias in aliases.items():
        if original in fields:
            extended_fields[alias] = fields[original]
    
    def replace_in_paragraph(paragraph, fields):
        for key, value in fields.items():
            placeholder = '{{' + key.upper() + '}}'
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value or ''))
    
    def replace_in_table(table, fields):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, fields)
                for nested_table in cell.tables:
                    replace_in_table(nested_table, fields)
    
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, extended_fields)
    
    for table in doc.tables:
        replace_in_table(table, extended_fields)
    
    return doc

def generate_invoice_pdf(fields):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    styles = getSampleStyleSheet()
    
    style_normal = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
    )
    style_bold = ParagraphStyle(
        'CustomBold',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        fontName='Helvetica-Bold',
    )
    style_center_bold = ParagraphStyle(
        'CustomCenterBold',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        fontName='Helvetica-Bold',
        alignment=TA_CENTER,
    )
    
    story = []
    
    # Логотип в верхнем левом углу
    logo_path = 'static:logo.png'
    if os.path.exists(logo_path):
        logo = Image(logo_path, width=1.5*cm, height=1.5*cm)
        logo_table = Table([[logo]], colWidths=[17*cm])
        logo_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(logo_table)
    story.append(Spacer(1, 0.5*cm))
    
    # Реквизиты продавца
    seller_info = [
        'Seller: Individual Entrepreneur HyperFlux',
        'IIN: 880923050176',
        'Address: 48, 21 Mamyr-7, Almaty, Kazakhstan, A10E9D2',
        'Website: https://wind4tune.com',
        'E-mail: support@wind4tune.com',
    ]
    for line in seller_info:
        story.append(Paragraph(line, style_normal))
    
    story.append(Spacer(1, 0.4*cm))
    
    # Банковские реквизиты продавца
    story.append(Paragraph('Seller Banking details:', style_bold))
    bank_info = [
        "Beneficiary's name: Individual Entrepreneur HyperFlux",
        'Bank name: Bank CenterCredit JSC',
        "Beneficiary's bank address: 38 Al-Farabi Ave., Almaty, A25D5G0, Republic of Kazakhstan",
        'SWIFT / BIC code: KCJBKZKX',
        f'IBAN / Account number: {fields.get("seller_iban", "KZ918562204233254882")}',
    ]
    for line in bank_info:
        story.append(Paragraph(line, style_normal))
    
    story.append(Spacer(1, 0.6*cm))
    
    # Номер и дата инвойса
    story.append(Paragraph(f'Invoice No.: {fields.get("invoice_number", "")}', style_normal))
    story.append(Paragraph(f'Date: {fields.get("invoice_date", "")}', style_normal))
    
    story.append(Spacer(1, 0.4*cm))
    
    # Данные клиента
    story.append(Paragraph(f'Customer: {fields.get("customer_name", "")}', style_normal))
    story.append(Paragraph(f'VAT number: {fields.get("customer_vat", "")}', style_normal))
    story.append(Paragraph(f'UEN: {fields.get("customer_uen", "")}', style_normal))
    story.append(Paragraph(f"Beneficiary's Address: {fields.get('customer_address', '')}", style_normal))
    
    story.append(Spacer(1, 0.4*cm))
    
    # Банк клиента
    story.append(Paragraph("Beneficiary's Bank:", style_normal))
    story.append(Paragraph(f'IBAN: {fields.get("customer_iban", "")}', style_normal))
    story.append(Paragraph(f'BIC: {fields.get("customer_bic", "")}', style_normal))
    story.append(Paragraph(f'Intermediary BIC: {fields.get("customer_intermediary_bic", "")}', style_normal))
    story.append(Paragraph(f'Bank address: {fields.get("customer_bank_address", "")}', style_normal))
    
    story.append(Spacer(1, 0.6*cm))
    
    # Текст обращения
    story.append(Paragraph('Dear Customer,', style_normal))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        "Below you will find your order's specifications and the commercial terms thereof. "
        "The order is regulated by the General Terms and Conditions of Sale available on our "
        "website at https://wind4tune.com. The purchase of an online event specified herein "
        "becomes binding only after receipt by the seller of the payment under this invoice in full.",
        style_normal
    ))
    
    story.append(Spacer(1, 0.5*cm))
    
    # Название мероприятия
    story.append(Paragraph(f'Virtual Event: {fields.get("event_name", "")}', style_center_bold))
    
    story.append(Spacer(1, 0.4*cm))
    
    # Таблица с деталями
    story.append(Paragraph('Additional Order Details', style_bold))
    story.append(Spacer(1, 0.2*cm))
    
    table_data = [
        ['Date:', fields.get("event_date", "")],
        ['Duration:', fields.get("event_time", "")],
        ['Additional participants:', fields.get("participants", "")],
        ['Additional price:', f'{fields.get("currency", "€")} {fields.get("amount", "")}'],
    ]
    
    table = Table(table_data, colWidths=[6*cm, 10*cm])
    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (0, -1), 0),
    ]))
    story.append(table)
    
    story.append(Spacer(1, 0.6*cm))
    
    # Финальный текст
    story.append(Paragraph(
        'By accepting and paying this invoice you warrant that you have the total legal '
        'capacity and authorisation to do so.',
        style_normal
    ))
    
    story.append(Spacer(1, 0.4*cm))
    
    # Подпись над чертой
    signature_path = 'static:signature.png'
    if os.path.exists(signature_path):
        sig_img = Image(signature_path, width=4*cm, height=1.5*cm)
        sig_table = Table(
            [
                [sig_img, ''],
                [Paragraph('______________________ / Sigalov Sergey', style_normal), '']
            ],
            colWidths=[8*cm, 9*cm]
        )
        sig_table.setStyle(TableStyle([
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
        ]))
        story.append(KeepTogether([sig_table]))
    else:
        story.append(Paragraph('______________________ / Sigalov Sergey', style_normal))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.json
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        doc_type = data.get('doc_type')
        fields = data.get('fields', {})
        
        if not doc_type:
            return jsonify({'error': 'doc_type is required'}), 400
        
        template_map = {
            'doc_contract': 'templates/contract_template.docx',
            'doc_appendix': 'templates/appendix_template.docx',
            'doc_invoice_ru': 'templates/invoice_ru_template.docx',
            'doc_act': 'templates/act_template.docx',
        }
        
        template_path = template_map.get(doc_type)
        if not template_path:
            return jsonify({'error': f'Unknown doc_type: {doc_type}'}), 400
        
        if not os.path.exists(template_path):
            return jsonify({'error': f'Template not found: {template_path}'}), 404
        
        doc = Document(template_path)
        doc = replace_placeholders(doc, fields)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        doc_number = fields.get('doc_number', 'draft')
        filename = f"{doc_type}_{doc_number}.docx"
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-pdf', methods=['POST'])
def generate_pdf():
    try:
        data = request.json
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        fields = data.get('fields', {})
        
        buffer = generate_invoice_pdf(fields)
        
        invoice_number = fields.get('invoice_number', 'draft')
        filename = f"invoice_{invoice_number}.pdf"
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
